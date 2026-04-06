from docx import Document

doc = Document(r'd:\PROJECTS\BloomingHives\assets\bloominghives-seo-content.docx')

with open(r'd:\PROJECTS\BloomingHives\seo_content.txt', 'w', encoding='utf-8') as f:
    for p in doc.paragraphs:
        f.write(p.text + '\n')

    # Also extract tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                f.write(cell.text + '\t')
            f.write('\n')

print("Done! Content written to seo_content.txt")
